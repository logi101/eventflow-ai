#!/usr/bin/env python3
"""
Readability metrics calculator.
Analyzes text complexity and provides readability scores.
"""

import argparse
import sys
from pathlib import Path

try:
    import textstat
    from docx import Document
except ImportError:
    print("Error: Required libraries not installed. Install with:")
    print("  pip install textstat python-docx")
    sys.exit(1)


class ReadabilityScorer:
    def __init__(self, language='en'):
        self.language = language

    def analyze_text(self, text):
        """Calculate readability metrics for text"""
        if not text.strip():
            return None

        scores = {
            'flesch_reading_ease': textstat.flesch_reading_ease(text),
            'flesch_kincaid_grade': textstat.flesch_kincaid_grade(text),
            'gunning_fog': textstat.gunning_fog(text),
            'smog_index': textstat.smog_index(text),
            'automated_readability_index': textstat.automated_readability_index(text),
            'coleman_liau_index': textstat.coleman_liau_index(text),
            'lexicon_count': textstat.lexicon_count(text, removepunct=True),
            'sentence_count': textstat.sentence_count(text),
            'avg_sentence_length': textstat.avg_sentence_length(text),
            'avg_syllables_per_word': textstat.avg_syllables_per_word(text),
            'difficult_words': textstat.difficult_words(text)
        }

        return scores

    def analyze_docx(self, file_path):
        """Analyze readability of DOCX file"""
        doc = Document(file_path)
        full_text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())

        # Overall document scores
        overall_scores = self.analyze_text(full_text)

        # Section-by-section analysis
        section_scores = []
        current_section = []
        current_heading = "Introduction"

        for para in doc.paragraphs:
            if not para.text.strip():
                continue

            # Check if it's a heading
            style_name = para.style.name if para.style else 'Normal'
            if 'Heading' in style_name or 'Title' in style_name:
                # Analyze previous section if exists
                if current_section:
                    section_text = '\n'.join(current_section)
                    scores = self.analyze_text(section_text)
                    if scores:
                        section_scores.append({
                            'heading': current_heading,
                            'scores': scores
                        })
                current_heading = para.text
                current_section = []
            else:
                current_section.append(para.text)

        # Analyze last section
        if current_section:
            section_text = '\n'.join(current_section)
            scores = self.analyze_text(section_text)
            if scores:
                section_scores.append({
                    'heading': current_heading,
                    'scores': scores
                })

        return overall_scores, section_scores

    def interpret_score(self, score_name, value):
        """Provide interpretation of readability scores"""
        interpretations = {
            'flesch_reading_ease': [
                (90, 100, "Very Easy (5th grade)"),
                (80, 90, "Easy (6th grade)"),
                (70, 80, "Fairly Easy (7th grade)"),
                (60, 70, "Standard (8th-9th grade)"),
                (50, 60, "Fairly Difficult (10th-12th grade)"),
                (30, 50, "Difficult (College level)"),
                (0, 30, "Very Difficult (College graduate)")
            ],
            'flesch_kincaid_grade': [
                (0, 6, "Elementary school"),
                (6, 9, "Middle school"),
                (9, 13, "High school"),
                (13, 16, "College"),
                (16, 100, "Graduate school")
            ]
        }

        if score_name in interpretations:
            for low, high, desc in interpretations[score_name]:
                if low <= value < high:
                    return desc
        return "N/A"

    def print_report(self, overall_scores, section_scores):
        """Print readability analysis report"""
        print("\n📈 Readability Analysis Report")
        print("=" * 60)
        print()

        if overall_scores:
            print("📊 Overall Document Scores:")
            print()
            print(f"  Flesch Reading Ease: {overall_scores['flesch_reading_ease']:.1f}")
            print(f"    → {self.interpret_score('flesch_reading_ease', overall_scores['flesch_reading_ease'])}")
            print()
            print(f"  Flesch-Kincaid Grade Level: {overall_scores['flesch_kincaid_grade']:.1f}")
            print(f"    → {self.interpret_score('flesch_kincaid_grade', overall_scores['flesch_kincaid_grade'])}")
            print()
            print(f"  Gunning Fog Index: {overall_scores['gunning_fog']:.1f}")
            print(f"  SMOG Index: {overall_scores['smog_index']:.1f}")
            print(f"  Automated Readability Index: {overall_scores['automated_readability_index']:.1f}")
            print(f"  Coleman-Liau Index: {overall_scores['coleman_liau_index']:.1f}")
            print()
            print(f"📝 Text Statistics:")
            print(f"  Total words: {overall_scores['lexicon_count']}")
            print(f"  Total sentences: {overall_scores['sentence_count']}")
            print(f"  Average sentence length: {overall_scores['avg_sentence_length']:.1f} words")
            print(f"  Average syllables per word: {overall_scores['avg_syllables_per_word']:.2f}")
            print(f"  Difficult words: {overall_scores['difficult_words']}")
            print()

        if section_scores:
            print("📖 Section-by-Section Analysis:")
            print()
            for section in section_scores:
                heading = section['heading'][:60] + '...' if len(section['heading']) > 60 else section['heading']
                scores = section['scores']
                print(f"  📄 {heading}")
                print(f"     Reading Ease: {scores['flesch_reading_ease']:.1f} | "
                      f"Grade Level: {scores['flesch_kincaid_grade']:.1f} | "
                      f"Words: {scores['lexicon_count']}")
                print()

    def get_suggestions(self, scores):
        """Generate improvement suggestions based on scores"""
        suggestions = []

        if scores['flesch_reading_ease'] < 60:
            suggestions.append("⚡ Consider simplifying sentences for better readability")

        if scores['avg_sentence_length'] > 25:
            suggestions.append("⚡ Average sentence length is high - try shorter sentences")

        if scores['difficult_words'] / scores['lexicon_count'] > 0.15:
            suggestions.append("⚡ High percentage of difficult words - consider simpler alternatives")

        if scores['flesch_kincaid_grade'] > 14:
            suggestions.append("⚡ Reading level is very high - may be difficult for general audience")

        return suggestions


def main():
    parser = argparse.ArgumentParser(description='Calculate readability metrics')
    parser.add_argument('file_path', help='Path to document file (.docx)')
    parser.add_argument('--language', '-l', choices=['en', 'he'],
                       default='en', help='Document language')

    args = parser.parse_args()

    if not Path(args.file_path).exists():
        print(f"Error: File not found: {args.file_path}")
        sys.exit(1)

    print(f"📈 Analyzing readability: {args.file_path}")
    if args.language == 'he':
        print("   Note: Readability metrics are optimized for English text")
    print()

    scorer = ReadabilityScorer(args.language)
    overall_scores, section_scores = scorer.analyze_docx(args.file_path)
    scorer.print_report(overall_scores, section_scores)

    # Print suggestions
    if overall_scores:
        suggestions = scorer.get_suggestions(overall_scores)
        if suggestions:
            print("💡 Improvement Suggestions:")
            for suggestion in suggestions:
                print(f"   {suggestion}")
            print()


if __name__ == '__main__':
    main()
